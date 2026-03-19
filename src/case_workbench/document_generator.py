#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
案件工作台 - 文书生成器
支持授权委托书、法定代表人身份证明书、答辩状的生成和Word导出
"""
import os
import re
from datetime import datetime
from typing import Dict, List, Optional, Any
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn


class DocumentGenerator:
    """文书生成器"""
    
    # 文书类型
    DOC_TYPE_POWER_OF_ATTORNEY = "power_of_attorney"      # 授权委托书
    DOC_TYPE_LEGAL_REP_CERT = "legal_rep_cert"            # 法定代表人身份证明书
    DOC_TYPE_DEFENSE_STATEMENT = "defense_statement"      # 答辩状
    
    def __init__(self, output_dir: Optional[str] = None):
        """
        初始化生成器
        
        Args:
            output_dir: 输出目录，默认为用户桌面/案件文书
        """
        if output_dir:
            self.output_dir = output_dir
        else:
            desktop = Path.home() / "Desktop"
            self.output_dir = desktop / "案件文书"
        
        self.output_dir = Path(self.output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def _set_chinese_font(self, run, font_name='宋体', font_size=12, bold=False):
        """设置中文字体"""
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(font_size)
        run.font.bold = bold
    
    def _create_document(self) -> Document:
        """创建标准Word文档"""
        doc = Document()
        
        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        style.font.size = Pt(12)
        
        return doc
    
    def generate_power_of_attorney(self, case_data: Dict[str, Any]) -> str:
        """
        生成授权委托书
        
        Args:
            case_data: 案件数据，包含以下字段：
                - plaintiff: 委托人（原告/申请人）
                - plaintiff_type: 委托人类型（自然人/法人）
                - defendant: 对方当事人（被告/被申请人）
                - case_cause: 案由
                - lawyer_name: 律师姓名
                - law_firm: 律师事务所
                - lawyer_license: 执业证号
                - court: 管辖法院
                
        Returns:
            生成的文件路径
        """
        doc = self._create_document()
        
        # 标题
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('授 权 委 托 书')
        self._set_chinese_font(run, font_name='黑体', font_size=22, bold=True)
        title.paragraph_format.space_after = Pt(30)
        
        # 委托人信息
        plaintiff = case_data.get('plaintiff', '')
        plaintiff_type = case_data.get('plaintiff_type', '自然人')
        
        p = doc.add_paragraph()
        run = p.add_run(f"委托人：{plaintiff}")
        self._set_chinese_font(run)
        
        # 添加委托人详细信息
        if plaintiff_type == '自然人':
            p = doc.add_paragraph()
            gender = case_data.get('plaintiff_gender', '______')
            id_card = case_data.get('plaintiff_id_card', '____________________')
            address = case_data.get('plaintiff_address', '________________________________')
            run = p.add_run(f"性别：{gender}  身份证号：{id_card}")
            self._set_chinese_font(run)
            
            p = doc.add_paragraph()
            run = p.add_run(f"住址：{address}")
            self._set_chinese_font(run)
        else:
            # 公司信息
            company_name = case_data.get('company_name', '')
            company_address = case_data.get('company_address', '________________________________')
            company_legal_rep = case_data.get('company_legal_rep', '____________')
            
            p = doc.add_paragraph()
            run = p.add_run(f"公司名称：{company_name}")
            self._set_chinese_font(run)
            
            p = doc.add_paragraph()
            run = p.add_run(f"住所地：{company_address}")
            self._set_chinese_font(run)
            
            p = doc.add_paragraph()
            run = p.add_run(f"法定代表人/负责人：{company_legal_rep}")
            self._set_chinese_font(run)
        
        doc.add_paragraph()  # 空行
        
        # 受委托人信息（办案律师）
        lawyer_name = case_data.get('lawyer_name', '')
        lawyer_gender = case_data.get('lawyer_gender', '')
        lawyer_id_card = case_data.get('lawyer_id_card', '')
        lawyer_phone = case_data.get('lawyer_phone', '')
        law_firm = case_data.get('law_firm', '')
        lawyer_license = case_data.get('lawyer_license', '')
        
        p = doc.add_paragraph()
        run = p.add_run(f"受委托人：{lawyer_name}")
        self._set_chinese_font(run)
        
        if lawyer_gender:
            p = doc.add_paragraph()
            run = p.add_run(f"性别：{lawyer_gender}")
            self._set_chinese_font(run)
        
        if lawyer_id_card:
            p = doc.add_paragraph()
            run = p.add_run(f"身份证号：{lawyer_id_card}")
            self._set_chinese_font(run)
        
        if lawyer_phone:
            p = doc.add_paragraph()
            run = p.add_run(f"联系电话：{lawyer_phone}")
            self._set_chinese_font(run)
        
        p = doc.add_paragraph()
        run = p.add_run(f"工作单位：{law_firm}")
        self._set_chinese_font(run)
        
        p = doc.add_paragraph()
        run = p.add_run(f"执业证号：{lawyer_license}")
        self._set_chinese_font(run)
        
        doc.add_paragraph()  # 空行
        
        # 委托事项
        defendant = case_data.get('defendant', '')
        case_cause = case_data.get('case_cause', '')
        
        p = doc.add_paragraph()
        content = f"现委托上列受委托人在我与{defendant}{case_cause}一案中，作为我的诉讼代理人。"
        run = p.add_run(content)
        self._set_chinese_font(run)
        
        doc.add_paragraph()  # 空行
        
        # 代理权限
        p = doc.add_paragraph()
        run = p.add_run("代理权限如下：")
        self._set_chinese_font(run)
        
        p = doc.add_paragraph()
        run = p.add_run("☑ 一般代理：代为起诉、应诉、代为陈述事实、参加辩论、代为调取证据、代为调解、代为签收法律文书等。")
        self._set_chinese_font(run)
        
        p = doc.add_paragraph()
        run = p.add_run("☐ 特别授权：代为承认、放弃、变更诉讼请求，进行和解，提起反诉或上诉。")
        self._set_chinese_font(run)
        
        doc.add_paragraph()  # 空行
        doc.add_paragraph()  # 空行
        
        # 签章区域
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run("委托人（签名/盖章）：______________")
        self._set_chinese_font(run)
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_str = datetime.now().strftime('%Y年%m月%d日')
        run = p.add_run(f"日期：{date_str}")
        self._set_chinese_font(run)
        
        # 保存文件
        filename = f"授权委托书_{plaintiff}_{datetime.now().strftime('%Y%m%d')}.docx"
        output_path = self.output_dir / filename
        doc.save(output_path)
        
        return str(output_path)
    
    def generate_legal_rep_cert(self, case_data: Dict[str, Any]) -> str:
        """
        生成法定代表人身份证明书
        
        Args:
            case_data: 案件数据，包含以下字段：
                - company_name: 公司名称
                - company_credit_code: 统一社会信用代码
                - company_address: 公司住所
                - company_legal_rep: 法定代表人姓名
                
        Returns:
            生成的文件路径
        """
        doc = self._create_document()
        
        # 标题
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('法定代表人身份证明书')
        self._set_chinese_font(run, font_name='黑体', font_size=22, bold=True)
        title.paragraph_format.space_after = Pt(30)
        
        # 使用公司信息字段
        company_name = case_data.get('company_name', '') or case_data.get('defendant', '')
        company_credit_code = case_data.get('company_credit_code', '')
        company_address = case_data.get('company_address', '')
        company_legal_rep = case_data.get('company_legal_rep', '')
        
        # 正文
        p = doc.add_paragraph()
        content = f"兹证明 {company_legal_rep} 在我单位担任 法定代表人/负责人 职务。"
        run = p.add_run(content)
        self._set_chinese_font(run)
        
        doc.add_paragraph()  # 空行
        
        # 单位信息
        p = doc.add_paragraph()
        run = p.add_run(f"单位名称：{company_name}")
        self._set_chinese_font(run)
        
        if company_credit_code:
            p = doc.add_paragraph()
            run = p.add_run(f"统一社会信用代码：{company_credit_code}")
            self._set_chinese_font(run)
        
        if company_address:
            p = doc.add_paragraph()
            run = p.add_run(f"单位住所：{company_address}")
            self._set_chinese_font(run)
        
        doc.add_paragraph()  # 空行
        doc.add_paragraph()  # 空行
        
        # 签章区域
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run("单位（盖章）：______________")
        self._set_chinese_font(run)
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_str = datetime.now().strftime('%Y年%m月%d日')
        run = p.add_run(f"日期：{date_str}")
        self._set_chinese_font(run)
        
        # 保存文件
        filename = f"法定代表人身份证明书_{company_name}_{datetime.now().strftime('%Y%m%d')}.docx"
        output_path = self.output_dir / filename
        doc.save(output_path)
        
        return str(output_path)
    
    def generate_defense_statement(self, case_data: Dict[str, Any],
                                   selected_facts: List[str],
                                   selected_defenses: List[str],
                                   selected_legal_basis: List[str],
                                   custom_content: Optional[Dict[str, str]] = None) -> str:
        """
        生成答辩状
        
        Args:
            case_data: 案件数据
            selected_facts: 选择的事实段落
            selected_defenses: 选择的抗辩要点
            selected_legal_basis: 选择的法律依据
            custom_content: 自定义内容，用于替换占位符
                
        Returns:
            生成的文件路径
        """
        doc = self._create_document()
        
        # 标题
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('民 事 答 辩 状')
        self._set_chinese_font(run, font_name='黑体', font_size=22, bold=True)
        title.paragraph_format.space_after = Pt(30)
        
        # 答辩人信息
        defendant = case_data.get('defendant', '')
        defendant_address = custom_content.get('defendant_address', '________________') if custom_content else '________________'
        legal_rep = custom_content.get('legal_rep', '________________') if custom_content else '________________'
        
        p = doc.add_paragraph()
        run = p.add_run(f"答辩人：{defendant}")
        self._set_chinese_font(run)
        
        p = doc.add_paragraph()
        run = p.add_run(f"住所地：{defendant_address}")
        self._set_chinese_font(run)
        
        p = doc.add_paragraph()
        run = p.add_run(f"法定代表人/负责人：{legal_rep}")
        self._set_chinese_font(run)
        
        doc.add_paragraph()  # 空行
        
        # 被答辩人信息
        plaintiff = case_data.get('plaintiff', '')
        
        p = doc.add_paragraph()
        run = p.add_run(f"被答辩人（原告）：{plaintiff}")
        self._set_chinese_font(run)
        
        doc.add_paragraph()  # 空行
        
        # 案由
        case_cause = case_data.get('case_cause', '')
        
        p = doc.add_paragraph()
        run = p.add_run(f"答辩人因与被答辩人{case_cause}一案，现提出答辩如下：")
        self._set_chinese_font(run)
        
        doc.add_paragraph()  # 空行
        
        # 事实部分
        if selected_facts:
            for i, fact in enumerate(selected_facts, 1):
                p = doc.add_paragraph()
                # 替换占位符
                fact_content = self._replace_placeholders(fact, case_data, custom_content)
                run = p.add_run(fact_content)
                self._set_chinese_font(run)
                doc.add_paragraph()  # 段落间空行
        
        # 抗辩要点
        if selected_defenses:
            for defense in selected_defenses:
                p = doc.add_paragraph()
                defense_content = self._replace_placeholders(defense, case_data, custom_content)
                run = p.add_run(defense_content)
                self._set_chinese_font(run)
                doc.add_paragraph()  # 段落间空行
        
        # 法律依据
        if selected_legal_basis:
            p = doc.add_paragraph()
            run = p.add_run("综上所述，本案应适用以下法律规定：")
            self._set_chinese_font(run)
            
            for legal in selected_legal_basis:
                p = doc.add_paragraph()
                run = p.add_run(f"• {legal}")
                self._set_chinese_font(run)
            
            doc.add_paragraph()  # 空行
        
        # 结语
        court = case_data.get('court', '贵院')
        
        p = doc.add_paragraph()
        run = p.add_run(f"综上所述，被告与原告之间{'不' if case_cause == '确认劳动关系纠纷' else ''}存在{case_cause.replace('纠纷', '')}，原告的诉讼请求缺乏事实和法律依据，恳请{court}依法驳回原告的全部诉讼请求。")
        self._set_chinese_font(run)
        
        doc.add_paragraph()  # 空行
        doc.add_paragraph()  # 空行
        
        # 尾部
        p = doc.add_paragraph()
        run = p.add_run("此致")
        self._set_chinese_font(run)
        
        p = doc.add_paragraph()
        run = p.add_run(court)
        self._set_chinese_font(run)
        
        doc.add_paragraph()  # 空行
        doc.add_paragraph()  # 空行
        
        # 签章
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"答辩人：{defendant}（盖章）")
        self._set_chinese_font(run)
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_str = datetime.now().strftime('%Y年%m月%d日')
        run = p.add_run(f"日期：{date_str}")
        self._set_chinese_font(run)
        
        # 保存文件
        filename = f"答辩状_{defendant}诉{plaintiff}_{datetime.now().strftime('%Y%m%d')}.docx"
        output_path = self.output_dir / filename
        doc.save(output_path)
        
        return str(output_path)
    
    def _replace_placeholders(self, content: str, case_data: Dict[str, Any],
                              custom_content: Optional[Dict[str, str]] = None) -> str:
        """替换内容中的占位符"""
        placeholders = {
            '{plaintiff}': case_data.get('plaintiff', ''),
            '{defendant}': case_data.get('defendant', ''),
            '{court}': case_data.get('court', ''),
            '{case_cause}': case_data.get('case_cause', ''),
            '{lawyer_name}': case_data.get('lawyer_name', ''),
            '{law_firm}': case_data.get('law_firm', ''),
            '{date}': datetime.now().strftime('%Y年%m月%d日'),
            '{third_party}': custom_content.get('third_party', '第三方公司') if custom_content else '第三方公司',
            '{contract_date}': custom_content.get('contract_date', '____年__月__日') if custom_content else '____年__月__日',
        }
        
        result = content
        for placeholder, value in placeholders.items():
            result = result.replace(placeholder, value)
        
        return result
    
    def preview_defense_statement(self, case_data: Dict[str, Any],
                                   selected_facts: List[str],
                                   selected_defenses: List[str],
                                   selected_legal_basis: List[str],
                                   custom_content: Optional[Dict[str, str]] = None) -> str:
        """
        预览答辩状内容（返回纯文本，用于界面预览）
        
        Returns:
            答辩状内容文本
        """
        lines = []
        
        # 标题
        lines.append("民 事 答 辩 状")
        lines.append("")
        
        # 当事人信息
        defendant = case_data.get('defendant', '')
        defendant_address = custom_content.get('defendant_address', '________________') if custom_content else '________________'
        legal_rep = custom_content.get('legal_rep', '________________') if custom_content else '________________'
        plaintiff = case_data.get('plaintiff', '')
        case_cause = case_data.get('case_cause', '')
        
        lines.append(f"答辩人：{defendant}")
        lines.append(f"住所地：{defendant_address}")
        lines.append(f"法定代表人/负责人：{legal_rep}")
        lines.append("")
        lines.append(f"被答辩人（原告）：{plaintiff}")
        lines.append("")
        lines.append(f"答辩人因与被答辩人{case_cause}一案，现提出答辩如下：")
        lines.append("")
        
        # 事实部分
        for fact in selected_facts:
            fact_content = self._replace_placeholders(fact, case_data, custom_content)
            lines.append(fact_content)
            lines.append("")
        
        # 抗辩要点
        for defense in selected_defenses:
            defense_content = self._replace_placeholders(defense, case_data, custom_content)
            lines.append(defense_content)
            lines.append("")
        
        # 法律依据
        if selected_legal_basis:
            lines.append("综上所述，本案应适用以下法律规定：")
            for legal in selected_legal_basis:
                lines.append(f"• {legal}")
            lines.append("")
        
        # 结语
        court = case_data.get('court', '贵院')
        lines.append(f"综上所述，被告与原告之间{'不' if case_cause == '确认劳动关系纠纷' else ''}存在{case_cause.replace('纠纷', '')}，原告的诉讼请求缺乏事实和法律依据，恳请{court}依法驳回原告的全部诉讼请求。")
        lines.append("")
        lines.append("此致")
        lines.append(court)
        lines.append("")
        lines.append("")
        lines.append(f"答辩人：{defendant}（盖章）")
        lines.append(f"日期：{datetime.now().strftime('%Y年%m月%d日')}")
        
        return '\n'.join(lines)
